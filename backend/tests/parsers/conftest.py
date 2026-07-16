"""Shared fixtures for parser tests."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from PIL import Image, ImageDraw


def build_text_pdf(text: str) -> bytes:
    """Build a minimal valid PDF containing one or more lines of text."""
    lines = text.splitlines() or [text]
    y = 720
    stream_parts: list[str] = []
    for line in lines:
        escaped = (
            line.replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)")
        )
        stream_parts.append(f"BT /F1 24 Tf 72 {y} Td ({escaped}) Tj ET")
        y -= 28
    stream = "\n".join(stream_parts)

    chunks: list[str] = []

    def add_chunk(content: str) -> None:
        chunks.append(content)

    add_chunk("1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
    add_chunk("2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n")
    add_chunk(
        "3 0 obj\n"
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        "/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\n"
        "endobj\n"
    )
    add_chunk(
        f"4 0 obj\n<< /Length {len(stream)} >>\nstream\n{stream}\nendstream\nendobj\n"
    )
    add_chunk("5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n")

    body = "".join(chunks)
    body_bytes = body.encode("latin-1")

    offsets: list[int] = [0]
    position = 0
    for chunk in chunks:
        offsets.append(position)
        position += len(chunk.encode("latin-1"))

    xref_lines = ["xref", f"0 {len(offsets)}", "0000000000 65535 f "]
    for offset in offsets[1:]:
        xref_lines.append(f"{offset:010d} 00000 n ")

    xref = "\n".join(xref_lines) + "\n"
    xref_start = len(body_bytes)
    trailer = f"trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{xref_start}\n%%EOF\n"

    return body_bytes + xref.encode("latin-1") + trailer.encode("latin-1")


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """Valid PDF with extractable text."""
    path = tmp_path / "resume.pdf"
    path.write_bytes(build_text_pdf("Jane Doe\nSoftware Engineer"))
    return path


@pytest.fixture
def empty_pdf(tmp_path: Path) -> Path:
    """PDF page with no text operators."""
    path = tmp_path / "empty.pdf"
    stream = ""
    chunks = [
        "1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        "2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        (
            "3 0 obj\n"
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            "/Contents 4 0 R >>\nendobj\n"
        ),
        f"4 0 obj\n<< /Length {len(stream)} >>\nstream\n{stream}\nendstream\nendobj\n",
    ]
    body = "".join(chunks).encode("latin-1")
    offsets = [0]
    position = 0
    for chunk in chunks:
        offsets.append(position)
        position += len(chunk.encode("latin-1"))
    xref_lines = ["xref", f"0 {len(offsets)}", "0000000000 65535 f "]
    for offset in offsets[1:]:
        xref_lines.append(f"{offset:010d} 00000 n ")
    xref = "\n".join(xref_lines) + "\n"
    xref_start = len(body)
    trailer = f"trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{xref_start}\n%%EOF\n"
    path.write_bytes(body + xref.encode("latin-1") + trailer.encode("latin-1"))
    return path


@pytest.fixture
def corrupted_pdf(tmp_path: Path) -> Path:
    """Invalid PDF bytes."""
    path = tmp_path / "corrupted.pdf"
    path.write_bytes(b"not-a-valid-pdf")
    return path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Valid DOCX with paragraph text."""
    path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Backend Developer")
    document.save(path)
    return path


@pytest.fixture
def empty_docx(tmp_path: Path) -> Path:
    """DOCX with no paragraph text."""
    path = tmp_path / "empty.docx"
    Document().save(path)
    return path


@pytest.fixture
def legacy_doc(tmp_path: Path) -> Path:
    """Legacy DOC file with embedded UTF-16-LE text."""
    path = tmp_path / "resume.doc"
    text = "Legacy Resume Content"
    payload = text.encode("utf-16-le")
    path.write_bytes(b"\x00" * 128 + payload + b"\x00" * 128)
    return path


@pytest.fixture
def sample_txt(tmp_path: Path) -> Path:
    """Plain text file with extra whitespace."""
    path = tmp_path / "job.txt"
    path.write_text("  Senior   Engineer  \r\n\r\n\r\nPython   FastAPI  \n", encoding="utf-8")
    return path


@pytest.fixture
def sample_png(tmp_path: Path) -> Path:
    """PNG image suitable for OCR."""
    path = tmp_path / "resume.png"
    image = Image.new("RGB", (400, 120), color="white")
    draw = ImageDraw.Draw(image)
    draw.text((20, 40), "OCR Sample Text", fill="black")
    image.save(path, format="PNG")
    return path


@pytest.fixture
def blank_png(tmp_path: Path) -> Path:
    """Solid white PNG image."""
    path = tmp_path / "blank.png"
    Image.new("RGB", (200, 200), color="white").save(path, format="PNG")
    return path


@pytest.fixture
def unsupported_file(tmp_path: Path) -> Path:
    """File with unsupported extension."""
    path = tmp_path / "data.csv"
    path.write_text("a,b,c", encoding="utf-8")
    return path
